"""
Document Processor Service

Handles extraction of text content from various document formats:
- PDF files
- Word documents (.docx)
- Plain text files
- Images (with OCR if available)
"""

import os
from pathlib import Path
from typing import Optional, Dict, Any
from django.core.files.uploadedfile import UploadedFile

# PDF processing
from PyPDF2 import PdfReader

# Word document processing  
from docx import Document as DocxDocument

# Image processing
from PIL import Image
import io


class DocumentProcessor:
    """
    Service for extracting text from various document formats.
    
    Supported formats:
    - PDF (.pdf)
    - Word (.docx)
    - Text (.txt)
    - Images (.png, .jpg, .jpeg) - basic support
    
    Usage:
        processor = DocumentProcessor()
        text = processor.extract_text(uploaded_file)
    """
    
    SUPPORTED_FORMATS = {
        'pdf': ['.pdf'],
        'word': ['.docx'],
        'text': ['.txt', '.md', '.rst'],
        'image': ['.png', '.jpg', '.jpeg', '.gif', '.webp'],
    }
    
    @classmethod
    def get_supported_extensions(cls) -> list:
        """Get all supported file extensions."""
        extensions = []
        for format_exts in cls.SUPPORTED_FORMATS.values():
            extensions.extend(format_exts)
        return extensions
    
    @classmethod
    def get_file_type(cls, filename: str) -> Optional[str]:
        """Determine file type from filename."""
        ext = Path(filename).suffix.lower()
        for file_type, extensions in cls.SUPPORTED_FORMATS.items():
            if ext in extensions:
                return file_type
        return None
    
    def extract_text(self, file: UploadedFile) -> Dict[str, Any]:
        """
        Extract text content from an uploaded file.
        
        Args:
            file: Django UploadedFile object
            
        Returns:
            Dictionary with 'text', 'file_type', 'page_count', 'success', 'error'
        """
        filename = file.name
        file_type = self.get_file_type(filename)
        
        if not file_type:
            return {
                'text': '',
                'file_type': None,
                'page_count': 0,
                'success': False,
                'error': f'Unsupported file format: {Path(filename).suffix}',
            }
        
        try:
            if file_type == 'pdf':
                return self._extract_from_pdf(file)
            elif file_type == 'word':
                return self._extract_from_docx(file)
            elif file_type == 'text':
                return self._extract_from_text(file)
            elif file_type == 'image':
                return self._extract_from_image(file)
        except Exception as e:
            return {
                'text': '',
                'file_type': file_type,
                'page_count': 0,
                'success': False,
                'error': str(e),
            }
    
    def _extract_from_pdf(self, file: UploadedFile) -> Dict[str, Any]:
        """Extract text from PDF file."""
        reader = PdfReader(file)
        
        text_parts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
        
        full_text = "\n\n".join(text_parts)
        
        return {
            'text': full_text,
            'file_type': 'pdf',
            'page_count': len(reader.pages),
            'success': True,
            'error': None,
        }
    
    def _extract_from_docx(self, file: UploadedFile) -> Dict[str, Any]:
        """Extract text from Word document."""
        doc = DocxDocument(file)
        
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        full_text = "\n\n".join(text_parts)
        
        return {
            'text': full_text,
            'file_type': 'word',
            'page_count': 1,  # Word doesn't have pages in the same way
            'success': True,
            'error': None,
        }
    
    def _extract_from_text(self, file: UploadedFile) -> Dict[str, Any]:
        """Extract text from plain text file."""
        content = file.read()
        
        # Try to decode with different encodings
        for encoding in ['utf-8', 'utf-16', 'latin-1', 'cp1252']:
            try:
                text = content.decode(encoding)
                break
            except UnicodeDecodeError:
                continue
        else:
            text = content.decode('utf-8', errors='replace')
        
        return {
            'text': text,
            'file_type': 'text',
            'page_count': 1,
            'success': True,
            'error': None,
        }
    
    def _extract_from_image(self, file: UploadedFile) -> Dict[str, Any]:
        """
        Handle image files.
        Note: Full OCR requires additional setup (Tesseract).
        For now, returns a message about the image.
        """
        try:
            image = Image.open(file)
            width, height = image.size
            
            # Basic image info without OCR
            return {
                'text': f'[Image: {file.name}, Size: {width}x{height}px]',
                'file_type': 'image',
                'page_count': 1,
                'success': True,
                'error': None,
                'requires_ocr': True,
            }
        except Exception as e:
            return {
                'text': '',
                'file_type': 'image',
                'page_count': 0,
                'success': False,
                'error': f'Failed to process image: {str(e)}',
            }
    
    def extract_text_from_path(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text from a file path (for already saved files).
        
        Args:
            file_path: Path to the file
            
        Returns:
            Same structure as extract_text()
        """
        path = Path(file_path)
        
        if not path.exists():
            return {
                'text': '',
                'file_type': None,
                'page_count': 0,
                'success': False,
                'error': f'File not found: {file_path}',
            }
        
        with open(path, 'rb') as f:
            # Create a simple object that mimics UploadedFile
            class FileWrapper:
                def __init__(self, file, name):
                    self._file = file
                    self.name = name
                    
                def read(self):
                    return self._file.read()
                
                def seek(self, pos):
                    return self._file.seek(pos)
            
            wrapper = FileWrapper(f, path.name)
            return self.extract_text(wrapper)
